# -*- coding: utf-8 -*-

# 将doc转为docx
from win32com import client as wc
import docx
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn

docName = u'D:\\pywork\\resource\\hah.docx'
word = wc.Dispatch('word.application')
doc = word.Documents.Open(u'D:\\pywork\\resource\\11.doc')
doc.SaveAs(docName, 12, False, "", True, "", False, False, False, False)
doc.Close()
word.Quit()

# 写入word
# 打开文档
document = docx.Document()
# 加入不同等级的标题
document.add_heading(u'MS WORD写入测试', 0)
document.add_heading(u'一级标题', 1)
document.add_heading(u'二级标题', 2)
# 添加文本
paragraph = document.add_paragraph(u'我们在做文本测试！')
# 设置字号
run = paragraph.add_run(u'设置字号、')
run.font.size = Pt(24)

# 设置字体
run = paragraph.add_run('Set Font,')
run.font.name = 'Consolas'

# 设置中文字体
run = paragraph.add_run(u'设置中文字体、')
run.font.name=u'宋体'
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

# 设置斜体
run = paragraph.add_run(u'斜体、')
run.italic = True

# 设置粗体
run = paragraph.add_run(u'粗体').bold = True

# 增加引用
document.add_paragraph('Intense quote', style='Intense Quote')

# 增加无序列表
document.add_paragraph(
    u'无序列表元素1', style='List Bullet'
)
document.add_paragraph(
    u'无序列表元素2', style='List Bullet'
)
# 增加有序列表
document.add_paragraph(
    u'有序列表元素1', style='List Number'
)
document.add_paragraph(
    u'有序列表元素2', style='List Number'
)
# 增加图像（此处用到图像image.bmp，请自行添加脚本所在目录中）
document.add_picture(u'D:\\pywork\\resource\\timg.jpg', width=Inches(1.25))

# 增加表格
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Name'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
# 再增加3行表格元素
for i in range(3):
    row_cells = table.add_row().cells
    row_cells[0].text = 'test'+str(i)
    row_cells[1].text = str(i)
    row_cells[2].text = 'desc'+str(i)

# 增加分页
document.add_page_break()

# 保存文件
document.save(u'D:\\pywork\\resource\\测试.docx')

# 读取word
document = docx.Document(docName)
line = [para.text for para in document.paragraphs]
for i in line:
    print(i)
